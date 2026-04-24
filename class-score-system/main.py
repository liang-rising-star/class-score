from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from jinja2 import Environment, FileSystemLoader
import sqlite3
import json
import hashlib
import datetime
import os
import io
import secrets
import string
import barcode
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

app = FastAPI()

# 获取当前文件所在目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 配置静态文件
static_dir = os.path.join(BASE_DIR, "static")
if os.path.exists(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")

# 提供logo.ico文件访问
@app.get("/logo.ico")
async def get_logo():
    logo_path = os.path.join(BASE_DIR, "logo.ico")
    if os.path.exists(logo_path):
        return FileResponse(logo_path, media_type="image/x-icon")
    raise HTTPException(status_code=404, detail="Logo not found")

# 配置Jinja2模板（直接使用，绕过starlette缓存的bug）
template_env = Environment(loader=FileSystemLoader(os.path.join(BASE_DIR, "templates")), autoescape=True)

def render_template(template_name: str, context: dict) -> HTMLResponse:
    template = template_env.get_template(template_name)
    html = template.render(**context)
    return HTMLResponse(html)

# 会话存储（简单内存存储，生产环境应使用Redis等）
sessions = {}

# 默认配置
def get_default_config():
    """返回默认配置，init_required=True，确保首次启动走管理员设置流程"""
    return {
        "system": {
            "name": "班级积分防伪核销管理系统",
            "init_required": True,
            "encryption_salt": secrets.token_hex(16)
        },
        "printing": {
            "page_size": "A4",
            "rows_per_page": 5,
            "columns_per_page": 2
        },
        "paths": {
            "database": "./data/db/scores.db",
            "output": "./data/output"
        },
        "admin": {
            "username": "",
            "password": ""
        }
    }

# 数据目录
DATA_DIR = os.path.join(BASE_DIR, "data")

# 配置文件路径
CONFIG_PATH = os.path.join(DATA_DIR, "config.json")

# 确保数据目录存在
def ensure_data_dir():
    os.makedirs(DATA_DIR, exist_ok=True)

# 读取配置文件（不存在则自动创建默认配置）
def read_config():
    ensure_data_dir()
    if not os.path.exists(CONFIG_PATH):
        # 配置文件不存在，创建默认配置并要求初始化管理员
        default_config = get_default_config()
        write_config(default_config)
        return default_config
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

# 写入配置文件
def write_config(config):
    ensure_data_dir()
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

# 获取数据库绝对路径
def get_db_path():
    config = read_config()
    db_path = config["paths"]["database"]
    if not os.path.isabs(db_path):
        db_path = os.path.join(BASE_DIR, db_path)
    return db_path

# 初始化数据库
def init_db():
    db_path = get_db_path()
    # 确保数据库目录存在
    db_dir = os.path.dirname(db_path)
    if db_dir:
        os.makedirs(db_dir, exist_ok=True)
    # 连接数据库
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    # 创建表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS scores (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        class_name TEXT NOT NULL,
        title TEXT NOT NULL,
        amount INTEGER NOT NULL,
        code TEXT UNIQUE NOT NULL,
        hash TEXT UNIQUE NOT NULL,
        issue_date TEXT NOT NULL,
        status TEXT DEFAULT '未使用'
    )''')
    conn.commit()
    conn.close()

# 生成积分码（8位短编号 + SHA-256哈希）
CODE_CHARS = string.ascii_uppercase + string.digits  # A-Z + 0-9

def generate_score_code(class_name, title, amount):
    """返回 (短编号, 哈希值)。短编号10位用于条码和核销，哈希用于防伪校验。"""
    short_code = ''.join(secrets.choice(CODE_CHARS) for _ in range(10))
    # 用短编号 + 业务数据生成哈希做防伪
    timestamp = datetime.datetime.now().isoformat()
    data = f"{short_code}{class_name}{title}{amount}{timestamp}"
    config = read_config()
    salt = config["system"]["encryption_salt"]
    hash_obj = hashlib.sha256((data + salt).encode())
    return short_code, hash_obj.hexdigest()

# 生成Code128条码图片到内存
def generate_barcode_image(barcode_data):
    """生成Code128条码PNG图片，返回BytesIO对象。barcode_data为条码内容（含面额的编码）"""
    Code128 = barcode.get_barcode_class('code128')
    code = Code128(barcode_data, writer=barcode.writer.ImageWriter())
    buf = io.BytesIO()
    code.write(buf, options={
        'module_width': 0.25,
        'module_height': 7.0,
        'write_text': False,
        'quiet_zone': 0.5,
        'format': 'PNG',
    })
    buf.seek(0)
    return buf

# 辅助：给段落添加 tab stop
def add_tab_stop(paragraph, tab_val, tab_pos):
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.makeelement(qn('w:tabs'), {})
    tab = tabs.makeelement(qn('w:tab'), {
        qn('w:val'): tab_val,
        qn('w:pos'): tab_pos,
    })
    tabs.append(tab)
    pPr.append(tabs)

# 辅助：生成面额文本框（锚定在段落右侧，三号字大字效果）
_shape_id_counter = [100]

def make_amount_textbox(amount_text):
    """生成面额文本框的 mc:AlternateContent XML 元素"""
    _shape_id_counter[0] += 1
    doc_pr_id = _shape_id_counter[0]
    font_sz = 32  # 三号 = 16pt = 半磅单位32
    cx = 706120   # 文本框宽
    cy = 325120   # 文本框高（约两行）
    pos_h = 1386840  # 水平偏移
    pos_v = 19050    # 垂直偏移

    xml_template = f'''<mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <mc:Choice Requires="wps">
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
        relativeHeight="251659264" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:posOffset>{pos_h}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>{pos_v}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="10160" b="10160"/>
        <wp:wrapNone/>
        <wp:docPr id="{doc_pr_id}" name="文本框 {doc_pr_id}"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBox="1"/>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{cx}" cy="{cy}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:solidFill><a:schemeClr val="lt1"/></a:solidFill>
                <a:ln w="6350"><a:noFill/></a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:pPr>
                      <w:rPr>
                        <w:sz w:val="{font_sz}"/>
                        <w:szCs w:val="{font_sz}"/>
                      </w:rPr>
                    </w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:b/>
                        <w:sz w:val="{font_sz}"/>
                        <w:szCs w:val="{font_sz}"/>
                      </w:rPr>
                      <w:t>{amount_text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                horzOverflow="overflow" vert="horz" wrap="square"
                lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                numCol="1" spcCol="0" rtlCol="0" fromWordArt="0"
                anchor="t" anchorCtr="0" forceAA="0" compatLnSpc="1">
                <a:noAutofit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </mc:Choice>
</mc:AlternateContent>'''
    return etree.fromstring(xml_template.encode('utf-8'))

# 导出Word文档
def export_to_word(class_name, title, amount, codes, issue_date):
    config = read_config()
    output_dir = config["paths"]["output"]
    if not os.path.isabs(output_dir):
        output_dir = os.path.join(BASE_DIR, output_dir)
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # 设置页面边距
    for section in doc.sections:
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    # 每页3列7行 = 21张卡片
    rows_per_page = 7
    columns = 3
    column_width = 6.0  # cm
    row_height_twip = '2131'  # ≈3.76cm，exact固定
    items_per_page = rows_per_page * columns  # 3列7行 = 21张卡片
    tab_right_pos = str(int(column_width * 567))

    for page_idx, page_codes in enumerate([codes[i:i+items_per_page] for i in range(0, len(codes), items_per_page)]):
        if page_idx > 0:
            doc.add_page_break()

        # 始终使用完整的8行表格，避免浪费纸张
        table = doc.add_table(rows=rows_per_page, cols=columns)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置列宽 + 固定行高 + 垂直居中
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = trPr.makeelement(qn('w:trHeight'), {
                qn('w:val'): row_height_twip,
                qn('w:hRule'): 'exact',
            })
            trPr.append(trHeight)

            for cell in row.cells:
                cell.width = Cm(column_width)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = tcPr.makeelement(qn('w:vAlign'), {qn('w:val'): 'center'})
                tcPr.append(vAlign)
                
                # 设置单元格边框为黑色
                tcBorders = tcPr.makeelement(qn('w:tcBorders'))
                
                # 上边框
                top = tcBorders.makeelement(qn('w:top'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',
                    qn('w:space'): '0',
                    qn('w:color'): '000000'
                })
                tcBorders.append(top)
                
                # 下边框
                bottom = tcBorders.makeelement(qn('w:bottom'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',
                    qn('w:space'): '0',
                    qn('w:color'): '000000'
                })
                tcBorders.append(bottom)
                
                # 左边框
                left = tcBorders.makeelement(qn('w:left'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',
                    qn('w:space'): '0',
                    qn('w:color'): '000000'
                })
                tcBorders.append(left)
                
                # 右边框
                right = tcBorders.makeelement(qn('w:right'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',
                    qn('w:space'): '0',
                    qn('w:color'): '000000'
                })
                tcBorders.append(right)
                
                tcPr.append(tcBorders)

        # 填充数据
        for i, code in enumerate(page_codes):
            row_idx = i // columns
            col_idx = i % columns
            cell = table.cell(row_idx, col_idx)

            # 清空默认段落
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)

            # 段落0：班级积分卡 + 右面板额文本框
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.first_line_indent = Pt(0)

            run_box = p.add_run()
            run_box._element.append(make_amount_textbox(f"{amount}分"))

            run = p.add_run("班级积分卡")
            run.bold = True
            run.font.size = Pt(10)

            # 段落1：班级（tab stop左靠）
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_tab_stop(p, 'right', tab_right_pos)
            run = p.add_run(f"班级：{class_name}")
            run.font.size = Pt(9)
            p.add_run('\t')

            # 段落2：标题 + 日期（tab stop左右锚定）
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            add_tab_stop(p, 'right', tab_right_pos)
            run = p.add_run(f"标题：{title}")
            run.font.size = Pt(9)
            p.add_run('\t')
            run = p.add_run(f"日期：{issue_date}")
            run.font.size = Pt(9)

            # 段落3：条码图片
            barcode_buf = generate_barcode_image(code)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(barcode_buf, width=Cm(5))

            # 段落4：编号
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(code)
            run.font.size = Pt(8)
            run.font.name = 'Consolas'
        
        # 对于最后一页，填充剩余的空白单元格
        total_cells = rows_per_page * columns
        for i in range(len(page_codes), total_cells):
            row_idx = i // columns
            col_idx = i % columns
            cell = table.cell(row_idx, col_idx)
            # 清空默认段落
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)

    filename = f"{class_name}_{title}_{amount}分_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath

# 检查是否已登录
def get_current_user(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id and session_id in sessions:
        return sessions[session_id]
    return None

# 安全防护中间件
@app.middleware("http")
async def security_middleware(request: Request, call_next):
    # 防止访问敏感文件，静默跳转登录页
    path = request.url.path
    sensitive_paths = ["/config.json", "/Dockerfile", "/docker-compose.yml", "/requirements.txt", "/db/", "/data/"]

    for sensitive_path in sensitive_paths:
        if sensitive_path in path:
            return RedirectResponse(url="/login")

    response = await call_next(request)
    return response

# 初始化管理员页面
@app.get("/", response_class=HTMLResponse)
async def init_admin(request: Request):
    config = read_config()
    if config["system"]["init_required"]:
        return render_template("init_admin.html", {"system_name": config["system"]["name"]})
    return RedirectResponse(url="/login")

# 提交管理员初始化
@app.post("/init_admin")
async def submit_init_admin(request: Request, username: str = Form(...), password: str = Form(...)):
    config = read_config()
    config["admin"]["username"] = username
    config["admin"]["password"] = password
    config["system"]["init_required"] = False
    write_config(config)
    return RedirectResponse(url="/login", status_code=303)

# 登录页面
@app.get("/login", response_class=HTMLResponse)
async def login(request: Request):
    config = read_config()
    if config["system"]["init_required"]:
        return RedirectResponse(url="/")
    return render_template("login.html", {"system_name": config["system"]["name"], "error": False})

# 提交登录
@app.post("/login")
async def submit_login(request: Request, username: str = Form(...), password: str = Form(...)):
    config = read_config()
    if config["admin"]["username"] != username or config["admin"]["password"] != password:
        return render_template("login.html", {"system_name": config["system"]["name"], "error": True})
    # 创建会话
    session_id = secrets.token_hex(32)
    sessions[session_id] = username
    response = RedirectResponse(url="/index", status_code=303)
    response.set_cookie(key="session_id", value=session_id, httponly=True)
    return response

# 登出
@app.get("/logout")
async def logout(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id and session_id in sessions:
        del sessions[session_id]
    response = RedirectResponse(url="/login")
    response.delete_cookie(key="session_id")
    return response

# 首页
@app.get("/index", response_class=HTMLResponse)
async def index(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    config = read_config()
    system_name = config["system"]["name"]
    return render_template("index.html", {"system_name": system_name})

# 生成积分页面
@app.get("/generate", response_class=HTMLResponse)
async def generate(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    config = read_config()
    system_name = config["system"]["name"]
    return render_template("generate.html", {"system_name": system_name, "success": False})

# 最近一次生成结果（用于POST后重定向展示）
last_generate_result = {}

# 提交生成积分
@app.post("/generate")
async def submit_generate(request: Request, class_name: str = Form(...), title: str = Form(...), amount: int = Form(...), quantity: int = Form(...)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    init_db()
    config = read_config()
    db_path = get_db_path()

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    codes = []
    issue_date = datetime.datetime.now().strftime("%Y-%m-%d")

    for _ in range(quantity):
        short_code, hash_val = generate_score_code(class_name, title, amount)
        codes.append(short_code)
        cursor.execute("SELECT * FROM scores WHERE code = ?", (short_code,))
        if not cursor.fetchone():
            cursor.execute(
                "INSERT INTO scores (class_name, title, amount, code, hash, issue_date, status) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (class_name, title, amount, short_code, hash_val, issue_date, "未使用")
            )

    conn.commit()
    conn.close()

    # 导出Word
    filepath = export_to_word(class_name, title, amount, codes, issue_date)
    filename = os.path.basename(filepath)

    # 存储结果，重定向到GET显示
    last_generate_result["data"] = {
        "system_name": config["system"]["name"],
        "filepath": filepath,
        "filename": filename,
        "quantity": quantity,
    }
    return RedirectResponse(url="/generate_result", status_code=303)

# 生成结果页（防止刷新重复提交）
@app.get("/generate_result", response_class=HTMLResponse)
async def generate_result(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    result = last_generate_result.pop("data", None)
    if not result:
        return RedirectResponse(url="/generate")
    return render_template("generate.html", {
        "system_name": result["system_name"],
        "success": True,
        "filepath": result["filepath"],
        "filename": result["filename"],
        "quantity": result["quantity"],
    })

# 核销页面
@app.get("/verify", response_class=HTMLResponse)
async def verify(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    config = read_config()
    system_name = config["system"]["name"]
    return render_template("verify.html", {"system_name": system_name, "results": None, "total_amount": 0, "error": None})

# 最近一次核销结果
last_verify_result = {}

# 提交核销
@app.post("/verify")
async def submit_verify(request: Request, codes: str = Form(None)):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    config = read_config()
    system_name = config["system"]["name"]

    # 空编码检查
    if not codes or not codes.strip():
        return render_template("verify.html", {
            "system_name": system_name,
            "results": None,
            "total_amount": 0,
            "error": "请输入或扫描积分编号后再核销",
        })

    init_db()
    db_path = get_db_path()

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    code_list = [code.strip().upper() for code in codes.split() if code.strip()]
    results = []
    total_amount = 0

    for i, code in enumerate(code_list, 1):
        cursor.execute("SELECT * FROM scores WHERE code = ? OR hash = ?", (code, code))
        score = cursor.fetchone()

        if not score:
            results.append({"index": i, "code": code, "amount": 0, "status": "错误", "message": "伪造编码"})
        elif score[7] == "已作废":
            results.append({"index": i, "code": code, "amount": score[3], "status": "错误", "message": "已核销编码"})
        else:
            cursor.execute("UPDATE scores SET status = '已作废' WHERE code = ?", (score[4],))
            results.append({"index": i, "code": code, "amount": score[3], "status": "成功", "message": "核销成功"})
            total_amount += score[3]

    conn.commit()
    conn.close()

    last_verify_result["data"] = {
        "system_name": system_name,
        "results": results,
        "total_amount": total_amount,
    }
    return RedirectResponse(url="/verify_result", status_code=303)

# 核销结果页（防止刷新重复提交）
@app.get("/verify_result", response_class=HTMLResponse)
async def verify_result(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    result = last_verify_result.pop("data", None)
    if not result:
        return RedirectResponse(url="/verify")
    return render_template("verify.html", {
        "system_name": result["system_name"],
        "results": result["results"],
        "total_amount": result["total_amount"],
        "error": None,
    })

# 最近一次设置结果
last_setting_result = None  # 'success' | 'wrong_password' | None

# 设置页面（GET，读取上次结果）
@app.get("/setting", response_class=HTMLResponse)
async def setting(request: Request):
    global last_setting_result
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login")
    config = read_config()
    system_name = config["system"]["name"]
    result = last_setting_result
    last_setting_result = None
    return render_template("setting.html", {
        "system_name": system_name,
        "success": result == "success",
        "wrong_password": result == "wrong_password",
    })

# 提交修改密码
@app.post("/setting")
async def submit_setting(request: Request, current_password: str = Form(...), password: str = Form(...)):
    global last_setting_result
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    config = read_config()
    if config["admin"]["password"] != current_password:
        last_setting_result = "wrong_password"
        return RedirectResponse(url="/setting", status_code=303)
    config["admin"]["password"] = password
    write_config(config)
    last_setting_result = "success"
    return RedirectResponse(url="/setting", status_code=303)

# 恢复出厂设置
def do_reset():
    """删除数据库、清空配置、删除输出文件"""
    # 删除数据库
    db_path = get_db_path()
    if os.path.exists(db_path):
        os.remove(db_path)
    # 删除output目录中的文件
    config = read_config()
    output_dir = config["paths"]["output"]
    if not os.path.isabs(output_dir):
        output_dir = os.path.join(BASE_DIR, output_dir)
    if os.path.isdir(output_dir):
        for f in os.listdir(output_dir):
            fp = os.path.join(output_dir, f)
            if os.path.isfile(fp):
                try:
                    os.remove(fp)
                except PermissionError:
                    pass  # 文件被占用，跳过
    # 重置配置（使用默认配置，init_required=True）
    write_config(get_default_config())
    # 清空会话
    sessions.clear()

# 提交恢复出厂设置
@app.post("/reset")
async def reset_system(request: Request, reset_password: str = Form(...)):
    global last_setting_result
    user = get_current_user(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    config = read_config()
    if config["admin"]["password"] != reset_password:
        last_setting_result = "wrong_password"
        return RedirectResponse(url="/setting", status_code=303)
    do_reset()
    return RedirectResponse(url="/", status_code=303)

# 下载生成的Word文件
@app.get("/download/{filename}")
async def download_file(request: Request, filename: str):
    if not get_current_user(request):
        return RedirectResponse(url="/login")
    config = read_config()
    output_dir = config["paths"]["output"]
    if not os.path.isabs(output_dir):
        output_dir = os.path.join(BASE_DIR, output_dir)
    filepath = os.path.join(output_dir, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(filepath, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 初始化数据库
init_db()
