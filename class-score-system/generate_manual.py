"""
生成用户手册 Word 文档
运行：python generate_manual.py
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = Pt(22)
style.paragraph_format.space_after = Pt(6)

for heading_level in [1, 2, 3]:
    hs = doc.styles[f'Heading {heading_level}']
    hs.font.name = '微软雅黑'
    hs._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_info_table(rows_data):
    """添加一个简洁的信息表"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, value) in enumerate(rows_data):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.width = Cm(4)
        cell_v.width = Cm(12)
        # key
        for p in cell_k.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell_k.add_paragraph()
        run = p.add_run(str(key))
        run.bold = True
        run.font.size = Pt(10)
        # value
        for p in cell_v.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell_v.add_paragraph()
        run = p.add_run(str(value))
        run.font.size = Pt(10)


def add_note(text):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("提示：")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def add_warning(text):
    """添加警告框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("警告：")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)


# ========== 正文内容 ==========

# 封面
add_title("班级积分防伪核销管理系统")
add_subtitle("用户操作手册")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run(f"文档版本：1.0\n生成日期：{datetime.date.today().strftime('%Y年%m月%d日')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ==================== 第一章 系统概述 ====================
doc.add_heading("一、系统概述", level=1)

doc.add_heading("1.1 系统简介", level=2)
doc.add_paragraph(
    "班级积分防伪核销管理系统是一款面向学校、培训机构等场景的积分管理工具。"
    "系统支持批量生成带有防伪条码的积分卡、导出Word文档打印，以及对积分卡进行扫码核销验证。"
    "每张积分卡包含唯一的10位编号和Code128条码，可有效防止伪造和重复使用。"
)

doc.add_heading("1.2 核心功能", level=2)
features = [
    ("积分卡生成", "批量生成积分卡，每张含唯一10位防伪编号 + Code128条码，导出Word文档直接打印"),
    ("积分卡核销", "输入或扫描积分编号进行核销验证，自动识别面额，显示核销结果和总面额"),
    ("防伪校验", "10位随机编号（36^10 ≈ 3.6万亿亿种组合）+ SHA-256哈希双重防伪"),
    ("管理功能", "管理员登录、密码修改、恢复出厂设置"),
]
for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f"● {title}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("1.3 系统环境", level=2)
add_info_table([
    ("运行环境", "Python 3.9+"),
    ("Web框架", "FastAPI + Uvicorn"),
    ("数据库", "SQLite 3（无需额外安装）"),
    ("访问地址", "http://localhost:8000"),
    ("支持浏览器", "Chrome / Edge / Firefox 等现代浏览器"),
    ("条码扫描", "支持任何可输出文本的Code128扫码器"),
])

doc.add_page_break()

# ==================== 第二章 安装与启动 ====================
doc.add_heading("二、安装与启动", level=1)

doc.add_heading("2.1 首次安装", level=2)
steps = [
    "确保已安装 Python 3.9 或更高版本",
    "打开命令行，进入项目目录：cd class-score-system",
    "安装依赖：pip install -r requirements.txt",
    "启动系统：python run.py",
    "浏览器打开 http://localhost:8000",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

add_note("依赖包括 fastapi、uvicorn、jinja2、python-docx、python-barcode、pillow、python-multipart")

doc.add_heading("2.2 Docker 部署", level=2)
doc.add_paragraph("如需使用 Docker 部署，执行以下命令：")
p = doc.add_paragraph()
run = p.add_run("docker-compose up -d")
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_note("Docker 部署时，数据库和导出文件会通过卷挂载持久化到宿主机的 ./db 和 ./output 目录。")

doc.add_heading("2.3 首次使用", level=2)
doc.add_paragraph(
    "首次打开系统时，会自动进入管理员初始化页面。在此页面设置管理员用户名和密码，"
    "设置完成后跳转到登录页面，使用刚设置的账号登录即可。"
)

doc.add_page_break()

# ==================== 第三章 功能使用 ====================
doc.add_heading("三、功能使用说明", level=1)

doc.add_heading("3.1 登录系统", level=2)
doc.add_paragraph("打开浏览器访问 http://localhost:8000，进入登录页面。输入管理员用户名和密码，点击「登录」按钮。")
doc.add_paragraph("如果输入错误，页面会提示「用户名或密码错误」，重新输入即可。")

doc.add_heading("3.2 系统首页", level=2)
doc.add_paragraph(
    "登录成功后进入系统首页，页面顶部显示系统名称，右上角有「设置」和「退出」链接。"
    "首页提供两个功能入口："
)
p = doc.add_paragraph()
run = p.add_run("● 生成积分")
run.bold = True
p.add_run(" — 批量生成积分卡并导出Word文档")
p = doc.add_paragraph()
run = p.add_run("● 核销积分")
run.bold = True
p.add_run(" — 扫码或输入编号验证核销")

doc.add_heading("3.3 生成积分卡", level=2)
doc.add_paragraph("点击首页的「生成积分」进入生成页面，填写以下信息：")

add_info_table([
    ("班级名称", "如：三年二班、八年十五班等"),
    ("积分卡标题", "如：优秀学生、期中奖励等"),
    ("单张积分面额", "每张积分卡的积分值（正整数）"),
    ("批量生成数量", "本次生成多少张积分卡"),
])

doc.add_paragraph("")
doc.add_paragraph("填写完成后点击「生成并导出」，系统将：")
steps = [
    "生成指定数量的积分卡，每张包含唯一的10位防伪编号",
    "将积分卡数据写入数据库",
    "自动生成Word文档（A4纸，每页10张卡片，2列×5行）",
    "页面显示成功提示和「下载Word文档」按钮",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

add_note("点击绿色的「下载Word文档」按钮可直接下载文件，无需到文件夹中查找。")

doc.add_heading("3.4 积分卡样式说明", level=2)
doc.add_paragraph("每张积分卡包含以下内容：")
card_layout = """
┌──────────────────┐
│    班级积分卡      │   ← 居中加粗标题
│ ─────────────── │   ← 分隔线
│ 班级：三年二班     │
│ 标题：优秀学生     │   ← 信息行
│ 面额：10分        │
│ 日期：2026-04-11   │
│                  │
│  ║║║║║║║║║║║║║  │   ← Code128条码（扫码可读出编号）
│   ABC123DEF4     │   ← 10位防伪编号
└──────────────────┘"""
p = doc.add_paragraph()
run = p.add_run(card_layout)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading("3.5 核销积分", level=2)
doc.add_paragraph("点击首页的「核销积分」进入核销页面，在文本框中输入要核销的积分编号：")
doc.add_paragraph("● 支持每行一个编号，也可用空格分隔多个编号")
doc.add_paragraph("● 可使用扫码枪扫描积分卡上的条码，编号会自动填入")
doc.add_paragraph("● 输入完成后点击「校验核销」")

doc.add_paragraph("")
doc.add_paragraph("核销结果会逐条显示：")
add_info_table([
    ("核销成功", "绿色提示，显示编号和面额"),
    ("伪造编码", "红色提示，该编号不存在于数据库中"),
    ("已核销编码", "红色提示，该编号已被核销过，不可重复使用"),
])
doc.add_paragraph("")
doc.add_paragraph("核销完成后，底部蓝色区域会显示「本次核销总面额：XX分」，仅统计成功核销的积分卡面额之和。")

add_note("面额信息存储在数据库中，核销时自动查询显示，无需在编号中编码面额。")

doc.add_heading("3.6 修改密码", level=2)
doc.add_paragraph(
    "点击页面右上角「设置」进入设置页面。在「新密码」输入框中输入新密码，"
    "点击「保存」即可。页面会提示「密码修改成功！」。"
)

add_warning("请妥善保管密码，系统当前未提供密码找回功能。如遗忘密码，需通过恢复出厂设置重置。")

doc.add_heading("3.7 退出登录", level=2)
doc.add_paragraph("点击页面右上角「退出」即可退出登录，系统会自动跳转回登录页面。")

doc.add_page_break()

# ==================== 第四章 恢复出厂设置 ====================
doc.add_heading("四、恢复出厂设置", level=1)

add_warning("恢复出厂设置将清除所有数据且不可恢复，请谨慎操作！")

doc.add_heading("4.1 网页端重置", level=2)
doc.add_paragraph("在设置页面底部「危险操作」区域，点击红色「恢复出厂设置」按钮，确认后系统将自动重置。")

doc.add_heading("4.2 双击批处理文件", level=2)
doc.add_paragraph("在项目文件夹中双击「恢复出厂设置.bat」文件，按提示操作即可。")

doc.add_heading("4.3 命令行重置", level=2)
doc.add_paragraph("打开命令行，进入项目目录，执行以下任一命令：")
p = doc.add_paragraph()
run = p.add_run("python run.py --reset")
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
run = p.add_run("python reset.py")
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading("4.4 重置会清除什么", level=2)
add_info_table([
    ("数据库 (db/scores.db)", "所有积分码和核销记录全部删除"),
    ("导出文件 (output/*.docx)", "所有已生成的Word文档"),
    ("管理员账号", "用户名和密码清空，需重新设置"),
    ("加密盐", "重新生成，旧积分卡将无法再核销"),
])
doc.add_paragraph("")
doc.add_paragraph("重置后，系统回到初始状态，访问时会显示管理员初始化页面，需重新设置账号。")

doc.add_heading("4.5 手动删除文件", level=2)
doc.add_paragraph("如果以上方式均不可用，可以手动删除以下文件达到重置效果：")
doc.add_paragraph("1. 删除 db/scores.db 文件（数据库）")
doc.add_paragraph("2. 删除 output 文件夹中的所有 .docx 文件")
doc.add_paragraph("3. 打开 config.json，将内容替换为：")
p = doc.add_paragraph()
config_text = '''{
  "system": {
    "name": "班级积分防伪核销管理系统",
    "init_required": true,
    "encryption_salt": "这里换成随机32位字符串"
  },
  "printing": { "page_size": "A4", "rows_per_page": 5, "columns_per_page": 2 },
  "paths": { "database": "./db/scores.db", "output": "./output" },
  "admin": { "username": "", "password": "" }
}'''
run = p.add_run(config_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ==================== 第五章 文件说明 ====================
doc.add_heading("五、项目文件说明", level=1)

add_info_table([
    ("main.py", "主程序，包含所有路由和业务逻辑"),
    ("run.py", "启动脚本，支持 --reset 参数"),
    ("reset.py", "恢复出厂设置脚本（可独立运行）"),
    ("恢复出厂设置.bat", "Windows 双击即执行的恢复出厂设置"),
    ("config.json", "系统配置文件（管理员账号、数据库路径等）"),
    ("requirements.txt", "Python 依赖清单"),
    ("Dockerfile", "Docker 镜像构建文件"),
    ("docker-compose.yml", "Docker Compose 编排文件"),
    ("db/scores.db", "SQLite 数据库（运行后自动生成）"),
    ("output/", "Word 文档输出目录"),
    ("templates/", "HTML 模板文件目录"),
    ("static/", "静态资源目录"),
])

doc.add_page_break()

# ==================== 第六章 数据库结构 ====================
doc.add_heading("六、数据库结构", level=1)

doc.add_paragraph("系统使用 SQLite 数据库，数据库文件位于 db/scores.db，包含一张 scores 表：")

# 表结构
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
headers = ["字段名", "类型", "约束", "说明"]
for j, h in enumerate(headers):
    cell = table.cell(0, j)
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)

rows = [
    ("id", "INTEGER", "主键自增", "记录ID"),
    ("class_name", "TEXT", "NOT NULL", "班级名称"),
    ("title", "TEXT", "NOT NULL", "积分卡标题"),
    ("amount", "INTEGER", "NOT NULL", "面额（分）"),
    ("code", "TEXT", "UNIQUE NOT NULL", "10位防伪短编号"),
    ("hash", "TEXT", "UNIQUE NOT NULL", "SHA-256哈希（64位）"),
    ("issue_date", "TEXT", "NOT NULL", "发放日期（YYYY-MM-DD）"),
]
for i, (name, typ, constraint, desc) in enumerate(rows):
    for j, val in enumerate([name, typ, constraint, desc]):
        cell = table.cell(i + 1, j)
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        run = p.add_run(val)
        run.font.size = Pt(9)
        if j == 0:
            run.font.name = 'Consolas'

doc.add_paragraph("")
doc.add_paragraph('status 字段（未列出，默认值「未使用」）：')
doc.add_paragraph("● 未使用 — 积分卡尚未核销，可正常使用")
doc.add_paragraph("● 已作废 — 积分卡已核销，不可重复使用")

doc.add_page_break()

# ==================== 第七章 常见问题 ====================
doc.add_heading("七、常见问题", level=1)

faqs = [
    ("Q：启动时提示端口被占用怎么办？",
     "A：修改 run.py 中的 port=8000 为其他端口号（如 8001），然后重新启动。"),
    ("Q：条码扫描器扫不出来怎么办？",
     "A：请确保打印清晰，条码区域不要模糊或被遮挡。建议使用至少600DPI打印。如果扫码器输出的是编号+额外字符，可在核销页面直接输入10位编号。"),
    ("Q：忘记了管理员密码怎么办？",
     "A：只能通过恢复出厂设置来重置（参见第四章）。重置后需重新设置管理员账号。"),
    ("Q：服务器重启后需要重新登录吗？",
     "A：是的。当前会话存储在内存中，服务器重启后会话失效，需重新登录。但数据库和配置不受影响。"),
    ("Q：可以同时在多台电脑上访问吗？",
     "A：可以。在局域网内，其他电脑可通过 http://服务器IP:8000 访问。但注意服务器重启后所有用户需重新登录。"),
    ("Q：恢复出厂设置后旧的积分卡还能用吗？",
     "A：不能。恢复出厂设置会重新生成加密盐，所有旧积分卡将无法核销验证。"),
    ("Q：生成的Word文档在哪里？",
     "A：在项目文件夹的 output 目录中。也可以在生成成功后直接点击网页上的「下载Word文档」按钮下载。"),
    ("Q：如何更换系统名称？",
     "A：用文本编辑器打开 config.json，修改 system.name 的值，重启系统即可生效。"),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)
    doc.add_paragraph("")

# ========== 保存 ==========
filepath = os.path.join(BASE_DIR, "用户手册.docx")
doc.save(filepath)
print(f"用户手册已生成: {filepath}")
